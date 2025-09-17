from copy import deepcopy
from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
import os, json
from docx.shared import Pt
from docx.oxml.ns import qn

from config import WORD_OUTPUT_PATH

class WordFactory:
    def __init__(self, json_path, word_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            self.codigos_adicionales = json.load(f)
            self.word_path = word_path

    def iter_block_items(self, parent):
        """
        Itera sobre los elementos (párrafos y tablas) de un documento o celda en el orden en el que aparecen.
        Se modificó para no incluir los sdt (content controls) ya que se eliminarán aparte.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        else:  # por ejemplo, una celda de tabla
            parent_elm = parent._tc
        for child in parent_elm.iterchildren():
            if child.tag.endswith('p'):
                yield Paragraph(child, parent)
            elif child.tag.endswith('tbl'):
                yield Table(child, parent)

    def extraer_secciones(self, doc):
        """
        Recorre el documento y agrupa de forma jerárquica en un diccionario utilizando tres niveles:
        - Nivel 1: párrafos con estilo "Heading 1"
        - Nivel 2: párrafos con estilo "Heading 2" (dentro del nivel 1)
        - Nivel 3: párrafos con estilo "Heading 3" (dentro del nivel 2)
        Los bloques (párrafos y tablas) que no son heading se agregan al último heading 3 detectado.
        
        La estructura resultante es del tipo:
        {
            "Texto de Heading 1": {
                "Texto de Heading 2": {
                    "Texto de Heading 3": [bloque1, bloque2, ...]
                },
                ...
            },
            ...
        }
        """
        secciones = {}
        current_h1 = None
        current_h2 = None
        current_h3 = None

        for bloque in self.iter_block_items(doc):
            if isinstance(bloque, Paragraph):
                texto = bloque.text.strip()
                estilo = str(bloque.style).lower()  # para comparar sin importar mayúsculas/minúsculas

                # Detectamos los headings según su nivel
                if "heading 1" in estilo:
                    current_h1 = texto
                    secciones[current_h1] = {}
                    current_h2 = None
                    current_h3 = None
                elif "heading 2" in estilo:
                    if current_h1 is None:
                        continue
                    current_h2 = texto
                    secciones[current_h1][current_h2] = {}
                    current_h3 = None
                elif "heading 3" in estilo:
                    if current_h1 is None or current_h2 is None:
                        continue
                    current_h3 = texto
                    # Se inicia la lista de bloques; se incluye el propio heading 3 en la lista
                    secciones[current_h1][current_h2][current_h3] = [bloque]
                else:
                    # Bloques que no son heading se agregan al último heading 3 detectado
                    if current_h1 is not None and current_h2 is not None and current_h3 is not None:
                        secciones[current_h1][current_h2][current_h3].append(bloque)
            else:
                if current_h1 is not None and current_h2 is not None and current_h3 is not None:
                    secciones[current_h1][current_h2][current_h3].append(bloque)
        return secciones

    def remove_block(self, block):
        """
        Elimina el bloque (párrafo o tabla) del documento utilizando su elemento XML.
        """
        block._element.getparent().remove(block._element)

    def remove_heading(self, doc, heading_text, heading_level):
        """
        Elimina del documento el párrafo con el heading especificado (texto y nivel)
        """
        for bloque in list(self.iter_block_items(doc)):
            if (isinstance(bloque, Paragraph) and 
                bloque.text.strip() == heading_text and 
                f"heading {heading_level}" in str(bloque.style).lower()):
                self.remove_block(bloque)

    def insert_paragraph_after(self, block, text, style=None):
        """
        Inserta un nuevo párrafo con el texto dado inmediatamente después del bloque indicado.
        Permite opcionalmente asignar un estilo (por ejemplo, 'Heading 3').
        Devuelve el objeto Paragraph insertado.
        """
        # Crear párrafo y texto
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_t = OxmlElement("w:t")
        new_t.text = text
        new_r.append(new_t)
        new_p.append(new_r)
        # Insertar en el documento tras el bloque dado
        block._element.addnext(new_p)

        # Envolver como objeto Paragraph
        new_paragraph = Paragraph(new_p, block._parent)
        if style is not None:
            new_paragraph.style = style

        # Aplicar formato de fuente y tamaño al run recién creado
        # Asumimos que solo existe un run en este párrafo recién insertado
        run = new_paragraph.runs[0]
        run.font.name = "Adif Fago No Regular"
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Adif Fago No Regular")

        return new_paragraph
    
    def count_elements(self,data):
        """
        Cuenta el número total de elementos en todas las listas que se encuentran
        en el nivel más profundo de un diccionario anidado.

        :param data: Diccionario anidado, donde los valores en el último nivel son listas.
        :return: Entero con el total de elementos en todas las listas.
        """
        total = 0

        def recurse(obj):
            nonlocal total
            # Si es diccionario, iterar por sus valores
            if isinstance(obj, dict):
                for value in obj.values():
                    recurse(value)
            # Si es lista, sumar su longitud
            elif isinstance(obj, list):
                total += len(obj)
            # En otros casos, ignorar o lanzar error según necesites
            else:
                pass

        recurse(data)
        return total

    def filter_sections(self, codigos_adicionales, ruta_entrada, ruta_salida = None):
        """
        Recorre la estructura jerárquica extraída con extraer_secciones y:
        - Elimina aquellas secciones (nivel 3) cuyo código (primeros 6 caracteres del heading) no estén en los códigos permitidos.
        - Si en algún caso un heading de nivel 2 queda sin secciones de nivel 3, se elimina ese heading.
        - Si un heading de nivel 1 queda sin secciones (nivel 2), se elimina también.
        - Para las secciones permitidas, se inserta al final:
          primero un párrafo con el título "UNIDADES" (con estilo "Heading 3")
          y luego otro párrafo con la información adicional.
        """
        # Calculamos los códigos permitidos a partir del valor de "CÓDIGO" en cada diccionario
        codigos_permitidos = {d["CÓDIGO"][:6] for d in codigos_adicionales}
        
        # Abrir el documento de entrada
        doc = Document(ruta_entrada)
        secciones = self.extraer_secciones(doc)

        total_number_of_sections = self.count_elements(secciones)
        print(f"Numero total de partidas de codigos a analizar: {total_number_of_sections}")

        
        # Recorrer la estructura jerárquica
        for h1 in list(secciones.keys()):
            for h2 in list(secciones[h1].keys()):
                for h3 in list(secciones[h1][h2].keys()):
                    codigo = h3[:6]
                    current_number_of_sections = self.count_elements(secciones)
                    print(f"Numero restante de partidas a analizar {total_number_of_sections - current_number_of_sections}/{total_number_of_sections}:")
    
                    if codigo not in codigos_permitidos:
                        # Eliminar todos los bloques asociados a este heading 3
                        for bloque in secciones[h1][h2][h3]:
                            self.remove_block(bloque)
                        del secciones[h1][h2][h3]
                    else:
                        # Sección permitida: insertar la información adicional
                        datos = [d for d in codigos_adicionales if d["CÓDIGO"][:6] == codigo]
                        datos_ordenados = sorted(codigos_adicionales, key=lambda d: d["CÓDIGO"])
                        if datos_ordenados:
                            ultimo_bloque = secciones[h1][h2][h3][-1]
                            p_unidades = self.insert_paragraph_after(ultimo_bloque, "UNIDADES", style="Heading 4")

                        texto_partidas = ""
                        for unidades in datos:
                            texto_partidas += f"CÓDIGO: {unidades['CÓDIGO']} | UD: {unidades['UD']} | RESUMEN: {unidades['RESUMEN']}\n"
                            # Insertar el párrafo con la información adicional justo después de "UNIDADES"
                        self.insert_paragraph_after(p_unidades, texto_partidas)
                # Si el nivel 2 quedó sin secciones (heading 3) válidas, eliminar el heading 2 del documento
                if not secciones[h1][h2]:
                    self.remove_heading(doc, h2, 2)
                    del secciones[h1][h2]
            # Si el nivel 1 quedó sin secciones (heading 2), eliminar el heading 1 del documento
            if not secciones[h1]:
                self.remove_heading(doc, h1, 1)
                del secciones[h1]
                
        # Guardar el documento modificado
        if ruta_salida:
            doc.save(ruta_salida)
        print(f"Documento procesado y guardado en {ruta_salida}")
        return secciones

    def concatenar_docs(self, doc_base, lista_rutas, ruta_salida):
        """
        Concatena en doc_base los contenidos de cada documento presente en lista_rutas.
        Se copian todos los bloques (párrafos, tablas, imágenes, etc.) utilizando deepcopy para preservar los objetos.
        """
        for ruta in lista_rutas:
            doc_aux = Document(ruta)
            for bloque in self.iter_block_items(doc_aux):
                doc_base.element.body.append(deepcopy(bloque._element))
        doc_base.save(ruta_salida)
        print(f"Documentos concatenados y guardados en {ruta_salida}")
        return doc_base

    def process_docx_files(self):
        """
        Procesa el archivo .docx proporcionado en self.word_path,
        aplicando filter_sections con los datos del JSON cargado en el constructor.
        Guarda el resultado en la misma ubicación con el sufijo '_procesado.docx'.
        
        :return: Ruta del archivo procesado si se procesó correctamente, None en caso contrario
        """
        if not os.path.exists(self.word_path):
            print(f"El archivo Word no existe: {self.word_path}")
            return None
            
        
        print(f"Procesando archivo: {self.word_path}")
        
        # Procesar el documento
        self.filter_sections(
            codigos_adicionales=self.codigos_adicionales, 
            ruta_entrada=self.word_path, 
            ruta_salida=WORD_OUTPUT_PATH
        )
        
        # Verificar si el documento procesado tiene contenido
        doc_temp = Document(WORD_OUTPUT_PATH)
        # Si hay al menos un bloque (párrafo o tabla), se considera que tiene contenido
        if any(True for _ in self.iter_block_items(doc_temp)):
            print(f"Archivo procesado correctamente: {WORD_OUTPUT_PATH}")
            return WORD_OUTPUT_PATH
        else:
            print(f"El archivo procesado está vacío: {WORD_OUTPUT_PATH}")
            return None
        
