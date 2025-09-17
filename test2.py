from docx import Document
import re
import os
import math


def is_heading3(paragraph):
    """Detecta si un párrafo tiene estilo Título 3 (Heading 3)."""
    return paragraph.style.name in ("Heading 3", "Título 3")


def is_heading12(paragraph):
    """Detecta si un párrafo tiene estilo Título 1 o 2 (Heading 1/2)."""
    return paragraph.style.name in (
        "Heading 1",
        "Título 1",
        "Heading 2",
        "Título 2",
    )


def safe_title(text, max_len=50):
    t = re.sub(r"\W+", "_", text).strip("_")
    return t[:max_len] if t else "SIN_TITULO"


def get_para_element_positions(doc):
    """
    Devuelve una lista con la posición (índice en body.iterchildren())
    de cada párrafo del documento, en el mismo orden que doc.paragraphs.
    También devuelve la lista completa de elementos (elems) y el body.
    """
    body = doc.element.body
    elems = list(body.iterchildren())
    id_map = {id(el): idx for idx, el in enumerate(elems)}
    positions = []
    for p in doc.paragraphs:
        pos = id_map.get(id(p._p))
        if pos is None:
            positions.append(-1)
        else:
            positions.append(pos)
    return positions, elems, body


def remove_outside_range(body, elems, keep_start, keep_end):
    """
    Elimina del body todos los elems cuyo índice no esté entre keep_start y keep_end (inclusive).
    """
    for idx in reversed(range(len(elems))):
        if idx < keep_start or idx > keep_end:
            body.remove(elems[idx])


def remove_heading1_2(doc):
    """
    Elimina del documento todos los párrafos que sean Heading 1 o Heading 2.
    Opera directamente sobre doc.element.body.
    """
    para_positions, elems, body = get_para_element_positions(doc)
    # recolectar índices de elems que corresponden a párrafos H1/H2
    indices_to_remove = []
    for i, p in enumerate(doc.paragraphs):
        if is_heading12(p):
            if i < len(para_positions):
                pos = para_positions[i]
                if pos != -1:
                    indices_to_remove.append(pos)
    # eliminar en orden inverso (y evitando duplicados)
    for idx in sorted(set(indices_to_remove), reverse=True):
        # protección: el índice podría estar fuera del rango si elems se ha modificado; comprobamos
        if 0 <= idx < len(elems):
            try:
                body.remove(elems[idx])
            except Exception:
                # en caso de algún fallo (no crítico), continuamos
                pass


def make_chunk_doc(original_path, keep_start_elem_idx, keep_end_elem_idx, out_path):
    """
    Crea un documento que contiene únicamente los elementos entre keep_start_elem_idx y keep_end_elem_idx
    (índices sobre elems, no sobre paragraphs). Después elimina Heading1/2.
    """
    doc = Document(original_path)
    # recalculamos elems/positions sobre esta copia
    para_positions, elems, body = get_para_element_positions(doc)
    remove_outside_range(body, elems, keep_start_elem_idx, keep_end_elem_idx)
    # eliminar Heading 1 y 2 si existen
    remove_heading1_2(doc)
    doc.save(out_path)


def split_chunk_into_sections(chunk_path, out_dir):
    """
    A partir de un chunk (documento más pequeño), genera los docx por sección Heading 3.
    Elimina Heading1/2 en cada sección también (por si quedara alguno).
    """
    chunk = Document(chunk_path)
    paras = chunk.paragraphs

    # encontrar índices de párrafos que son Heading3 dentro del chunk
    section_indices = []
    for idx, p in enumerate(paras):
        if is_heading3(p):
            title_clean = safe_title(p.text)
            section_indices.append((idx, title_clean))
    # añadir final
    section_indices.append((len(paras), "FIN"))

    for i in range(len(section_indices) - 1):
        start_para, title = section_indices[i]
        end_para, _ = section_indices[i + 1]

        # Reabrir chunk original (más pequeño) para recortar a la sección
        doc = Document(chunk_path)
        para_positions, elems, body = get_para_element_positions(doc)

        # sanity check
        if start_para >= len(para_positions) or (end_para - 1) >= len(para_positions):
            keep_start = 0
            keep_end = len(elems) - 1
        else:
            keep_start = para_positions[start_para]
            keep_end = para_positions[end_para - 1]

        # eliminar todo lo que no pertenezca
        remove_outside_range(body, elems, keep_start, keep_end)

        # eliminar Heading1/2 que pudieran quedar dentro de la sección
        remove_heading1_2(doc)

        # construir nombre: solo el código extraído (antes del primer guion bajo si lo hay)
        codigo = title.split("_")[0] if "_" in title else title
        filename = f"{codigo}.docx"
        out_path = os.path.join(out_dir, filename)
        doc.save(out_path)
        print(f"  Guardado sección: {out_path}")


def split_doc_by_heading3_parallel(input_path, output_dir, n_chunks=10):
    """
    Proceso en dos fases:
     1) Divide el documento original en n_chunks sub-documentos (por número de Heading3).
     2) Recorre cada chunk y genera los docx finales por sección.
    En la creación de cada chunk y en cada sección se eliminan Heading 1 y 2.
    """
    if n_chunks < 1:
        n_chunks = 1

    original = Document(input_path)
    paras = original.paragraphs

    # Detectar índices de inicio de cada Heading3 en el original
    section_starts = []
    section_titles = []
    for idx, p in enumerate(paras):
        if is_heading3(p):
            section_starts.append(idx)
            section_titles.append(safe_title(p.text))
    if not section_starts:
        raise ValueError("No se han encontrado Heading 3 en el documento.")

    # añadir índice final (len(paras)) para facilitar cálculos
    section_starts_with_end = section_starts[:] + [len(paras)]

    # determinar agrupamiento en chunks (distribución aproximada)
    total_sections = len(section_starts)
    per_chunk = math.ceil(total_sections / n_chunks)
    # ajustamos n_chunks real si hay menos secciones que chunks pedidas
    n_chunks_real = math.ceil(total_sections / per_chunk)

    # crear carpetas
    os.makedirs(output_dir, exist_ok=True)
    chunks_dir = os.path.join(output_dir, "chunks")
    os.makedirs(chunks_dir, exist_ok=True)
    sections_dir = os.path.join(output_dir, "sections")
    os.makedirs(sections_dir, exist_ok=True)

    print(
        f"Total de secciones: {total_sections}. Creando {n_chunks_real} chunks (aprox {per_chunk} secciones por chunk)."
    )

    # generar chunks
    for chunk_i in range(n_chunks_real):
        # secciones cubiertas por este chunk (índices en section_starts)
        s_idx = chunk_i * per_chunk
        e_idx = min((chunk_i + 1) * per_chunk - 1, total_sections - 1)

        start_para = section_starts_with_end[s_idx]
        # si e_idx + 1 < len(section_starts_with_end) entonces el siguiente inicio marca el final,
        # en caso contrario final es len(paras)
        if (e_idx + 1) < len(section_starts_with_end):
            end_para = section_starts_with_end[e_idx + 1]
        else:
            end_para = len(paras)

        # Para crear el chunk calculamos los índices sobre elems usando una copia del original
        doc = Document(input_path)
        para_positions, elems, body = get_para_element_positions(doc)

        # convertir start_para/end_para (índices sobre paragraphs del original) a índices sobre elems
        keep_start = para_positions[start_para]
        keep_end = (
            para_positions[end_para - 1]
            if (end_para - 1) < len(para_positions)
            else len(elems) - 1
        )

        chunk_filename = f"chunk_{chunk_i + 1:02d}.docx"
        chunk_path = os.path.join(chunks_dir, chunk_filename)

        # eliminar fuera del rango y luego Heading1/2
        remove_outside_range(body, elems, keep_start, keep_end)
        # El chunk ya no tendrá Heading1/2
        remove_heading1_2(doc)
        doc.save(chunk_path)
        print(
            f"Creado chunk {chunk_i + 1}/{n_chunks_real}: {chunk_path} (secciones {s_idx + 1} a {e_idx + 1})"
        )

    # ahora recorremos los chunks y los dividimos en secciones finales
    chunk_files = sorted(os.listdir(chunks_dir))
    for cf in chunk_files:
        chunk_path = os.path.join(chunks_dir, cf)
        print(f"Procesando {chunk_path} ...")
        split_chunk_into_sections(chunk_path, sections_dir)

    print("Proceso completado.")
    print(f"Archivos de secciones guardados en: {sections_dir}")


if __name__ == "__main__":
    INPUT_DOCX = "PPT9010_completo_limpio.docx"  # ruta a tu documento original
    OUTPUT_DIR = "secciones2"  # Carpeta de salida
    N_CHUNKS = 15  # ajusta a lo que quieras (p.ej. 10)

    split_doc_by_heading3_parallel(INPUT_DOCX, OUTPUT_DIR, n_chunks=N_CHUNKS)
