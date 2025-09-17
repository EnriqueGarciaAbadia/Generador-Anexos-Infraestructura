import os
import glob
from docx import Document
from docxcompose.composer import Composer

class Test3Factory:
    def __init__(self, original_docx, sections_dir, id_list, output_docx):
        self.original_docx = original_docx
        self.sections_dir = sections_dir
        self.id_list = id_list
        self.output_docx = output_docx

    def find_section_file(self, section_dir, identifier):
        import re
        match = re.search(r"[A-Z]{3}\d{3}", identifier)
        matches = []
        short_id = identifier
        if match:
            short_id = match.group(0)
            pattern = os.path.join(section_dir, f"*{short_id}*.docx")
            matches = glob.glob(pattern)
        if not matches:
            print(f"Aviso: No se encontró ningún archivo para el identificador '{identifier}' (patrón usado: '{short_id}')")
            return None
        return matches[0]

    def update_heading3_title(self, doc_path, identifier, index):
        """
        Crea una copia temporal del documento doc_path, modifica el título de nivel 3 (Heading 3)
        sustituyendo el identificador recortado (3 letras + 3 números) por el identificador completo,
        elimina el símbolo '$' si existe, añade un número incremental al principio del título,
        y mantiene el formato original: fuente "Adif Fago No Regular", subrayado, negrita y tamaño 11.
        Devuelve la ruta del nuevo archivo temporal.
        """
        import re
        import tempfile
        import shutil
        from docx import Document
        from docx.shared import Pt

        temp_fd, temp_path = tempfile.mkstemp(suffix=".docx")
        os.close(temp_fd)
        shutil.copy2(doc_path, temp_path)

        match = re.search(r"[A-Z]{3}\d{3}", identifier)
        if match:
            short_id = match.group(0)
        else:
            short_id = identifier

        doc = Document(temp_path)
        for p in doc.paragraphs:
            if p.style.name in ("Heading 3", "Título 3"):
                clean_text = p.text.replace("$", "")
                new_title = re.sub(re.escape(short_id), identifier, clean_text, count=1)
                final_title = f"III.{index:02d} {new_title} "

                p.clear()
                run = p.add_run(final_title)
                run.font.name = "Adif Fago No Regular"
                run.bold = True
                run.underline = True
                run.italic = False
                run.font.size = Pt(11)
                r = run._element
                rPr = r.get_or_add_rPr()
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), "Adif Fago No Regular")
                rFonts.set(qn('w:hAnsi'), "Adif Fago No Regular")
                rFonts.set(qn('w:eastAsia'), "Adif Fago No Regular")
                rFonts.set(qn('w:cs'), "Adif Fago No Regular")
                rPr.append(rFonts)
                break
        doc.save(temp_path)
        return temp_path

    def merge_sections_with_composer(self):
        import tempfile
        base_doc = Document(self.original_docx)
        composer = Composer(base_doc)
        temp_files = []
        codigos_no_añadidos = []
        idx = 1
        for ident in self.id_list:
            path = self.find_section_file(self.sections_dir, ident)
            if path is None:
                print(f"⚠️ Se omite el identificador '{ident}' porque no se encontró archivo.")
                codigos_no_añadidos.append(ident)
                continue
            temp_path = self.update_heading3_title(path, ident, idx)
            temp_files.append(temp_path)
            print(f"⟳ Concatenando sección '{ident}' desde: {temp_path}")
            subdoc = Document(temp_path)
            composer.append(subdoc)
            idx += 1
        composer.save(self.output_docx)
        print(f"✅ Documento final guardado en: {self.output_docx}")
        for temp_path in temp_files:
            os.remove(temp_path)

        # Guardar los códigos no añadidos en ficheros/codigos_no_añadidos.txt
        txt_path = os.path.join("ficheros", "codigos_no_añadidos.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            for codigo in codigos_no_añadidos:
                f.write(f"{codigo}\n")
        print(f"Archivo de códigos no añadidos guardado en: {txt_path}")

# Ejemplo de uso:
# factory = Test3Factory(
#     original_docx="original.docx",
#     sections_dir="secciones2",
#     id_list=[...],
#     output_docx="new_doc.docx"
# )
# factory.merge_sections_with_composer()
